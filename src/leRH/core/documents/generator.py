from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fpdf import FPDF
from openai import OpenAI

from leRH.config import settings
from leRH.db.models import Job, User

logger = logging.getLogger(__name__)

DATA_DIR = Path(__file__).resolve().parent.parent.parent.parent.parent / "data"
GENERATED_DIR = DATA_DIR / "generated"

_FONT_PATHS = [
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
]
_BOLD_FONT_PATHS = [
    "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf",
]
_ITALIC_FONT_PATHS = [
    "/usr/share/fonts/dejavu/DejaVuSans-Oblique.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
]

FONT_PATH = next((p for p in _FONT_PATHS if Path(p).exists()), None)
BOLD_FONT_PATH = next((p for p in _BOLD_FONT_PATHS if Path(p).exists()), None)
ITALIC_FONT_PATH = next((p for p in _ITALIC_FONT_PATHS if Path(p).exists()), None)
PDF_FONT = "DejaVu" if FONT_PATH and BOLD_FONT_PATH else "Helvetica"

CV_GENERATION_PROMPT = """Tu es un expert en rédaction de CV ATS-optimisés pour le marché
ouest-africain (Togo, Bénin, Côte d'Ivoire, Sénégal). Tu maîtrises les normes des recruteurs
internationaux (ONG, multinationales, banques) ET locaux (PME, administrations).

═══════════════════════════════════════════════
RÈGLES DE RÉDACTION STRICTES — NE PAS DÉROGER
═══════════════════════════════════════════════

RÉSUMÉ PROFESSIONNEL (obligatoire, 3-4 phrases max) :
- Phrase 1 : Objectif vers le titre exact du poste visé + années d'expérience uniquement si elles
  sont présentes dans le profil candidat
- Phrase 2 : 2-3 compétences techniques clés tirées du profil candidat, en privilégiant celles qui
  correspondent à l'offre d'emploi
- Phrase 3 : 1 réalisation quantifiée uniquement si le chiffre est présent dans le profil candidat
- Phrase 4 : Proposition de valeur unique du candidat pour CE poste

COMPÉTENCES CLÉS :
- 8 à 12 compétences maximum, 3 mots maximum par compétence
- Mélanger : hard skills (70%) + soft skills (30%)
- Reprendre les mots-clés de l'offre uniquement s'ils correspondent à des compétences du profil
- Ordre : de la plus pertinente à la moins pertinente pour CE poste

EXPÉRIENCES PROFESSIONNELLES :
- Ordre anti-chronologique STRICT (la plus récente en premier)
- Pour chaque poste : 3 à 5 bullets MAXIMUM
- Chaque bullet = 1 ACTION + 1 CONTEXTE/RÉSULTAT présent dans le profil
- Verbes d'action interdits (trop vagues) : "géré", "réalisé", "effectué", "participé"
- Verbes d'action IMPOSÉS : optimisé, déployé, dirigé, conçu, négocié, réduit, augmenté,
  automatisé, supervisé, développé, mis en place, coordonné, piloté, structuré
- Si pas de chiffres disponibles dans le profil : ne pas inventer de chiffres, utiliser un contexte
  qualitatif présent dans le profil
- Limiter à 10 ans d'historique maximum (ou 3 postes max)

FORMATION :
- Ordre anti-chronologique
- Mentionner la mention/spécialisation si disponible
- Inclure uniquement les formations certifiantes explicitement listées dans le profil candidat
- NE PAS mettre le baccalauréat si le candidat a un diplôme supérieur

LANGUES & CERTIFICATIONS :
- Format standardisé : "Langue (Niveau CECRL)" ex: "Français (C2 - Natif)", "Anglais (B2 - Courant)"
- Pour les certifications : "Nom cert. — Organisme émetteur — Année"
- Ne jamais transformer une compétence (AWS, Kubernetes, Azure, etc.) en certification.
- Si aucune certification n'est listée dans la section CERTIFICATIONS du profil, retourne "certifications": [].

═══════════════════════════════════════════
PROFIL DU CANDIDAT
═══════════════════════════════════════════
{profile}

═══════════════════════════════════════════
OFFRE D'EMPLOI CIBLE
═══════════════════════════════════════════
{job}

═══════════════════════════════════════════
INSTRUCTIONS FINALES
═══════════════════════════════════════════
1. Analyse d'abord les MOTS-CLÉS de l'offre (skills, diplômes requis, verbes d'action utilisés)
2. Assure-toi que ces mots-clés apparaissent NATURELLEMENT dans le résumé et les compétences
3. Tu peux reformuler les informations du profil, mais pas créer de certifications, diplômes, dates,
   entreprises, postes, outils maîtrisés ou résultats chiffrés absents du profil.
4. NE JAMAIS inventer des dates d'emploi, entreprises, postes, langues, compétences, diplômes,
   certifications, clients, volumes, pourcentages, nombres d'utilisateurs ou résultats mesurables.
5. Tout le texte doit être en FRANÇAIS impeccable (zéro faute orthographique)

Réponds UNIQUEMENT avec ce JSON valide (aucun texte avant ou après) :
{{
  "summary": "...",
  "core_competencies": ["...", "..."],
  "experience": [
    {{
      "company": "Nom exact de l'entreprise",
      "location": "Ville, Pays",
      "title": "Titre exact du poste",
      "start_date": "MM/YYYY",
      "end_date": "MM/YYYY ou Présent",
      "bullets": [
        "Verbe d'action fort + contexte précis + résultat/impact mesurable",
        "..."
      ]
    }}
  ],
  "education": [
    {{
      "degree": "Intitulé exact du diplôme + spécialisation",
      "institution": "Nom complet de l'établissement",
      "year": "YYYY ou En cours"
    }}
  ],
  "certifications": ["Nom cert. — Organisme — Année"],
  "languages": ["Langue (Niveau CECRL - Qualification)"]
}}"""

COVER_LETTER_PROMPT = """Tu es un expert en rédaction de lettres de motivation pour le marché
de l'emploi ouest-africain. Tu maîtrises les codes de communication professionnelle
au Togo, au Bénin et dans les pays de l'UEMOA.

═══════════════════════════════════════════════
STRUCTURE IMPOSÉE — 4 PARAGRAPHES PRÉCIS
═══════════════════════════════════════════════

PARAGRAPHE 1 — ACCROCHE (4-5 lignes max) :
Objectif : capter l'attention du recruteur en 5 secondes.
- Ouvrir avec UNE réalisation concrète et chiffrée du candidat (la plus impressionnante)
  Ex: "Fort de 5 ans d'expérience en développement web ayant permis de réduire de 40%
  le temps de chargement des applications d'une PME togolaise..."
- Puis exprimer l'intérêt SPÉCIFIQUE pour CE poste et CETTE entreprise
  (utiliser des éléments concrets tirés de l'offre : nom de l'entreprise, secteur, mission)
- NE PAS commencer par "Je me permets de..." ou "Suite à votre annonce..."
- NE PAS commencer par "Je" (règle typographique française)

PARAGRAPHE 2 — ADÉQUATION COMPÉTENCES/POSTE (6-8 lignes) :
Objectif : démontrer que le candidat EST le profil recherché.
- Sélectionner les 3 exigences PRINCIPALES de l'offre d'emploi
- Pour chacune, citer une preuve CONCRÈTE tirée du profil du candidat
- Format implicite : "Le poste exige X → J'ai démontré X en faisant Y (résultat Z)"
- Utiliser des connecteurs logiques : "En effet,", "Par ailleurs,", "De surcroît,"
- Intégrer 2-3 mots-clés techniques de l'offre de manière naturelle

PARAGRAPHE 3 — MOTIVATION ET VALEUR AJOUTÉE (4-6 lignes) :
Objectif : montrer pourquoi CETTE entreprise et pas une autre.
- Citer un élément spécifique qui attire le candidat (secteur, mission, valeurs, projets)
- Expliquer en quoi le candidat apportera quelque chose de NOUVEAU/DIFFÉRENT
- Projeter dans le futur : "Je suis convaincu de pouvoir contribuer à..."
- Évoquer la dimension culturelle si pertinente (contexte ouest-africain, langues locales, etc.)

PARAGRAPHE 4 — APPEL À L'ACTION (3-4 lignes) :
Objectif : obtenir un entretien.
- Exprimer la disponibilité pour un entretien de manière proactive (pas passive)
  Ex: "Je serais ravi de vous présenter plus en détail..." (PAS "Dans l'attente...")
- Mentionner la disponibilité concrète si connue
- Formule de politesse adaptée au contexte professionnel africain
  (respectueuse mais pas servile)

═══════════════════════════════════════════
RÈGLES STYLISTIQUES STRICTES
═══════════════════════════════════════════
- Longueur : 300-380 mots MAXIMUM (une page A4 avec marges standard)
- Ton : professionnel, chaleureux, confiant (jamais arrogant ni servile)
- Vocabulaire : riche mais accessible, zéro jargon incompréhensible
- Éviter absolument : "dynamique", "rigoureux", "motivé" sans preuve concrète
- Chaque affirmation doit être soit prouvée soit illustrée

═══════════════════════════════════════════
PROFIL DU CANDIDAT
═══════════════════════════════════════════
{profile}

═══════════════════════════════════════════
OFFRE D'EMPLOI CIBLE
═══════════════════════════════════════════
{job}

ANALYSE DE CORRESPONDANCE DISPONIBLE :
{match_analysis}

Réponds UNIQUEMENT avec ce JSON valide :
{{
  "recipient": "À l'attention de [Titre] [Nom si connu] / Direction des Ressources Humaines",
  "subject": "Candidature au poste de [Titre exact] — [Nom du candidat]",
  "body_paragraphs": [
    "Paragraphe 1 (accroche forte, PAS commencer par Je)...",
    "Paragraphe 2 (adéquation compétences/poste)...",
    "Paragraphe 3 (motivation et valeur ajoutée)...",
    "Paragraphe 4 (appel à l'action proactif)..."
  ],
  "closing": "Je vous prie d'agréer, [Titre], l'expression de mes salutations distinguées."
}}"""


@dataclass
class GeneratedCV:
    summary: str
    core_competencies: list[str]
    experience: list[dict]
    education: list[dict]
    certifications: list[str]
    languages: list[str]


@dataclass
class GeneratedCoverLetter:
    recipient: str
    subject: str
    body_paragraphs: list[str]
    closing: str


class DocumentGenerationError(Exception):
    pass


class DocumentGenerator:
    def __init__(self) -> None:
        self._client = OpenAI(
            api_key=settings.openai_api_key.get_secret_value(),
            base_url=settings.openai_base_url,
            timeout=settings.openai_timeout,
        )
        GENERATED_DIR.mkdir(parents=True, exist_ok=True)

    def _sanitize_filename(self, text: str) -> str:
        """Supprime les accents et caractères non-ASCII pour éviter les bugs WhatsApp/URL."""
        import re
        import unicodedata

        # Normalisation NFKD pour séparer les caractères de base des accents
        normalized = unicodedata.normalize("NFKD", text)
        # On garde seulement les caractères ASCII (ce qui supprime les accents)
        ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
        # Remplacement des espaces par des underscores
        no_spaces = ascii_text.replace(" ", "_")
        # Nettoyage strict : ne garder que l'alphanumérique et les underscores
        clean_name = re.sub(r"[^a-zA-Z0-9_]", "", no_spaces)
        return clean_name

    def _build_profile_text(self, user: User, cv_analysis: dict | None = None) -> str:
        """Construit un texte de profil structuré avec sections délimitées.

        Utilise des marqueurs de section explicites pour que le LLM comprenne
        la hiérarchie et l'importance relative des informations.

        Priorité des données :
        1. Champs directs du User (nom, pays, activité…)
        2. Profil structuré extrait du vrai CV via `cv_analysis` (table cvs.analysis)
        3. Texte brut extrait du CV si disponible
        """
        sections = []

        # ── IDENTITÉ ──────────────────────────────────────────────────
        identity_parts = [f"Nom complet : {user.name}"]
        if user.activity:
            identity_parts.append(f"Titre professionnel actuel : {user.activity}")
        if user.city:
            identity_parts.append(f"Ville : {user.city}")
        if user.country:
            identity_parts.append(f"Pays : {user.country}")
        if user.phone:
            identity_parts.append(f"Téléphone : {user.phone}")
        if user.email:
            identity_parts.append(f"Email : {user.email}")
        if user.linkedin_url:
            identity_parts.append(f"LinkedIn : {user.linkedin_url}")
        if user.github_url:
            identity_parts.append(f"GitHub : {user.github_url}")
        sections.append("── IDENTITÉ ──\n" + "\n".join(identity_parts))

        # ── RÉSUMÉ PROFESSIONNEL (depuis analyse CV) ───────────────────
        cv_profile: dict = {}
        if cv_analysis and isinstance(cv_analysis, dict):
            cv_profile = cv_analysis.get("profile", {})

        if cv_analysis and cv_analysis.get("analysis"):
            sections.append(
                "── RÉSUMÉ PROFESSIONNEL (extrait du CV réel uploadé) ──\n"
                + cv_analysis["analysis"]
            )

        if user.summary_override:
            sections.append(
                "── RÉSUMÉ PERSONNALISÉ (défini par le candidat — prioritaire) ──\n"
                + user.summary_override
            )

        # ── DIPLÔME & FORMATION ────────────────────────────────────────
        diploma = cv_profile.get("diploma") or user.diploma
        edu_section = []
        if user.educations:
            for edu in user.educations:
                line = f"• {edu.degree}"
                if edu.field:
                    line += f" en {edu.field}"
                if edu.institution:
                    line += f" — {edu.institution}"
                if edu.year:
                    line += f" ({edu.year})"
                edu_section.append(line)
        elif cv_profile.get("education"):
            for edu in cv_profile["education"]:
                if isinstance(edu, dict):
                    line = f"• {edu.get('degree', '')} en {edu.get('field', '')} — {edu.get('institution', '')} ({edu.get('year', '')})"
                    edu_section.append(line)
        elif diploma:
            edu_section.append(f"• Diplôme : {diploma}")

        if edu_section:
            sections.append("── FORMATION (ordre anti-chronologique) ──\n" + "\n".join(edu_section))

        # ── EXPÉRIENCES PROFESSIONNELLES ───────────────────────────────
        exp_section = []
        if user.experiences:
            exps = sorted(user.experiences, key=lambda e: e.start_date or "0000", reverse=True)
            for exp in exps:
                header = f"• {exp.title} @ {exp.company}"
                if exp.location:
                    header += f" ({exp.location})"
                dates = ""
                if exp.start_date or exp.end_date:
                    dates = f" | {exp.start_date or '?'} → {exp.end_date or 'Présent'}"
                exp_section.append(header + dates)
                if exp.description:
                    for line in exp.description.split("\n"):
                        if line.strip():
                            exp_section.append(f"  → {line.strip()}")
        elif cv_profile.get("experiences"):
            for exp in cv_profile["experiences"]:
                if isinstance(exp, dict):
                    header = f"• {exp.get('title', '')} @ {exp.get('company', '')}"
                    if exp.get("location"):
                        header += f" ({exp['location']})"
                    if exp.get("start_date") or exp.get("end_date"):
                        header += (
                            f" | {exp.get('start_date', '?')} → {exp.get('end_date', 'Présent')}"
                        )
                    exp_section.append(header)
                    if exp.get("description"):
                        exp_section.append(f"  → {exp['description']}")
        elif user.experience:
            exp_section.append(user.experience)

        if exp_section:
            sections.append("── EXPÉRIENCES PROFESSIONNELLES ──\n" + "\n".join(exp_section))

        # ── COMPÉTENCES ────────────────────────────────────────────────
        skills_from_cv: list[str] = []
        if cv_profile.get("skills"):
            raw = cv_profile["skills"]
            skills_from_cv = raw if isinstance(raw, list) else [str(raw)]

        skills_from_user: list[str] = []
        if user.skills:
            s = user.skills
            if isinstance(s, list):
                skills_from_user = [str(x) for x in s]
            elif isinstance(s, dict):
                skills_from_user = list(s.values())

        all_skills = list(dict.fromkeys(skills_from_cv + skills_from_user))
        if all_skills:
            skills_lines = []
            for i in range(0, len(all_skills), 4):
                skills_lines.append(" | ".join(all_skills[i : i + 4]))
            sections.append("── COMPÉTENCES ──\n" + "\n".join(skills_lines))

        # ── CERTIFICATIONS ─────────────────────────────────────────────
        if cv_profile.get("certifications"):
            certs = cv_profile["certifications"]
            if isinstance(certs, list) and certs:
                sections.append("── CERTIFICATIONS ──\n" + "\n".join(f"• {c}" for c in certs))

        # ── LANGUES ────────────────────────────────────────────────────
        langs_from_cv: list = cv_profile.get("languages", [])
        langs_source = langs_from_cv if langs_from_cv else (user.languages or [])
        if langs_source and isinstance(langs_source, list):
            lang_strs = []
            for lang in langs_source:
                if isinstance(lang, dict):
                    lang_strs.append(
                        f"• {lang.get('language', '')} — {lang.get('level', 'Niveau non précisé')}"
                    )
                else:
                    lang_strs.append(f"• {lang}")
            sections.append("── LANGUES ──\n" + "\n".join(lang_strs))

        # ── INFORMATIONS SUPPLÉMENTAIRES ───────────────────────────────
        extra = []
        if user.availability:
            extra.append(f"Disponibilité : {user.availability}")
        if user.website_url:
            extra.append(f"Portfolio/Site : {user.website_url}")
        if user.address:
            extra.append(f"Adresse : {user.address}")
        if extra:
            sections.append("── INFORMATIONS SUPPLÉMENTAIRES ──\n" + "\n".join(extra))

        return "\n\n".join(sections)

    @staticmethod
    def _normalized_token_set(text: str) -> set[str]:
        words = re.findall(r"[a-zA-ZÀ-ÿ0-9]+", text.lower())
        return {word for word in words if len(word) > 2}

    @classmethod
    def _certification_matches_source(cls, generated: str, source: str) -> bool:
        generated_tokens = cls._normalized_token_set(generated)
        source_tokens = cls._normalized_token_set(source)
        if not generated_tokens or not source_tokens:
            return False

        overlap = generated_tokens & source_tokens
        short_side = min(len(generated_tokens), len(source_tokens))
        return len(overlap) >= min(2, short_side)

    @classmethod
    def _extract_allowed_certifications(cls, cv_analysis: dict | None) -> list[str]:
        if not cv_analysis or not isinstance(cv_analysis, dict):
            return []

        profile = cv_analysis.get("profile", {})
        if not isinstance(profile, dict):
            return []

        raw_certs = profile.get("certifications", [])
        if not isinstance(raw_certs, list):
            raw_certs = [raw_certs]

        certs = [str(cert).strip() for cert in raw_certs if str(cert).strip()]

        for edu in profile.get("education", []) or []:
            if not isinstance(edu, dict):
                continue
            line = " ".join(
                str(edu.get(key) or "") for key in ("degree", "field", "institution", "year")
            ).strip()
            if re.search(
                r"\b(certification|certifié|certifiee|certifiante|formation)\b", line, re.I
            ):
                certs.append(line)

        return list(dict.fromkeys(certs))

    @classmethod
    def _filter_certifications(
        cls, generated_certs: list[str], cv_analysis: dict | None
    ) -> list[str]:
        allowed_certs = cls._extract_allowed_certifications(cv_analysis)
        if not allowed_certs:
            return []

        filtered = []
        for cert in generated_certs:
            cert_text = str(cert).strip()
            if cert_text and any(
                cls._certification_matches_source(cert_text, source) for source in allowed_certs
            ):
                filtered.append(cert_text)

        return list(dict.fromkeys(filtered))

    @staticmethod
    def _profile_from_analysis(cv_analysis: dict | None) -> dict:
        if not cv_analysis or not isinstance(cv_analysis, dict):
            return {}
        profile = cv_analysis.get("profile", {})
        return profile if isinstance(profile, dict) else {}

    @staticmethod
    def _source_text_from_analysis(cv_analysis: dict | None, fallback: str = "") -> str:
        if not cv_analysis or not isinstance(cv_analysis, dict):
            return fallback

        raw_text = cv_analysis.get("source_text") or cv_analysis.get("extracted_text")
        if raw_text:
            return str(raw_text)

        return fallback + "\n" + json.dumps(cv_analysis, ensure_ascii=False)

    @classmethod
    def _supported_by_source(cls, text: str, source_text: str) -> bool:
        tokens = cls._normalized_token_set(text) - {
            "avec",
            "dans",
            "pour",
            "sur",
            "des",
            "les",
            "une",
            "aux",
            "togo",
            "lome",
            "présent",
            "present",
        }
        if not tokens:
            return False

        source_tokens = cls._normalized_token_set(source_text)
        overlap = tokens & source_tokens
        if len(tokens) == 1:
            return bool(overlap)
        return len(overlap) >= min(2, len(tokens))

    @staticmethod
    def _number_tokens(text: str) -> set[str]:
        return set(re.findall(r"\d+(?:[.,]\d+)?", text))

    @classmethod
    def _has_unsupported_numbers(cls, text: str, source_text: str) -> bool:
        numbers = cls._number_tokens(text)
        if not numbers:
            return False
        source_numbers = cls._number_tokens(source_text)
        return any(number not in source_numbers for number in numbers)

    @classmethod
    def _remove_unsupported_numbered_sentences(cls, text: str, source_text: str) -> str:
        sentences = re.split(r"(?<=[.!?])\s+", text.strip())
        kept = [
            sentence
            for sentence in sentences
            if sentence and not cls._has_unsupported_numbers(sentence, source_text)
        ]
        return " ".join(kept)

    @staticmethod
    def _source_experiences(user: User, cv_analysis: dict | None) -> list[dict]:
        if user.experiences:
            return [
                {
                    "company": exp.company,
                    "location": exp.location,
                    "title": exp.title,
                    "start_date": exp.start_date,
                    "end_date": exp.end_date,
                    "description": exp.description,
                }
                for exp in user.experiences
            ]

        experiences = DocumentGenerator._profile_from_analysis(cv_analysis).get("experiences", [])
        return [exp for exp in experiences if isinstance(exp, dict)]

    @staticmethod
    def _source_educations(user: User, cv_analysis: dict | None) -> list[dict]:
        if user.educations:
            return [
                {
                    "degree": edu.degree,
                    "institution": edu.institution,
                    "field": edu.field,
                    "year": edu.year,
                }
                for edu in user.educations
            ]

        educations = DocumentGenerator._profile_from_analysis(cv_analysis).get("education", [])
        return [edu for edu in educations if isinstance(edu, dict)]

    @staticmethod
    def _source_languages(user: User, cv_analysis: dict | None) -> list:
        profile = DocumentGenerator._profile_from_analysis(cv_analysis)
        languages = profile.get("languages") or user.languages or []
        return languages if isinstance(languages, list) else []

    @classmethod
    def _best_source_match(
        cls, generated: dict, sources: list[dict], fields: tuple[str, ...]
    ) -> dict | None:
        generated_text = " ".join(str(generated.get(field) or "") for field in fields)
        generated_tokens = cls._normalized_token_set(generated_text)
        if not generated_tokens:
            return None

        best_source = None
        best_score = 0
        for source in sources:
            source_text = " ".join(str(source.get(field) or "") for field in fields)
            score = len(generated_tokens & cls._normalized_token_set(source_text))
            if score > best_score:
                best_score = score
                best_source = source

        return best_source if best_score >= 1 else None

    @classmethod
    def _sanitize_experience(
        cls, generated: list[dict], user: User, cv_analysis: dict | None, source_text: str
    ) -> list[dict]:
        sources = cls._source_experiences(user, cv_analysis)
        sanitized = []

        for exp in generated:
            if not isinstance(exp, dict):
                continue
            source = cls._best_source_match(exp, sources, ("company", "title"))
            if not source:
                continue

            bullets = []
            for bullet in exp.get("bullets", []) or []:
                bullet_text = str(bullet).strip()
                if bullet_text and not cls._has_unsupported_numbers(bullet_text, source_text):
                    bullets.append(bullet_text)

            if not bullets and source.get("description"):
                bullets = [
                    line.strip() for line in str(source["description"]).split("\n") if line.strip()
                ][:3]

            sanitized.append(
                {
                    "company": source.get("company") or exp.get("company", ""),
                    "location": source.get("location") or exp.get("location", ""),
                    "title": source.get("title") or exp.get("title", ""),
                    "start_date": source.get("start_date") or "",
                    "end_date": source.get("end_date") or "",
                    "bullets": bullets[:5],
                }
            )

        return sanitized

    @classmethod
    def _sanitize_education(
        cls, generated: list[dict], user: User, cv_analysis: dict | None
    ) -> list[dict]:
        sources = cls._source_educations(user, cv_analysis)
        sanitized = []

        for edu in generated:
            if not isinstance(edu, dict):
                continue
            source = cls._best_source_match(edu, sources, ("degree", "institution", "field"))
            if not source:
                continue

            degree_parts = [str(source.get("degree") or "").strip()]
            if source.get("field") and str(source["field"]).lower() not in degree_parts[0].lower():
                degree_parts.append(str(source["field"]).strip())

            sanitized.append(
                {
                    "degree": " en ".join(part for part in degree_parts if part),
                    "institution": source.get("institution") or edu.get("institution", ""),
                    "year": source.get("year") or "",
                }
            )

        return sanitized

    @classmethod
    def _sanitize_languages(
        cls, generated: list[str], user: User, cv_analysis: dict | None
    ) -> list[str]:
        sources = cls._source_languages(user, cv_analysis)
        sanitized = []

        for lang in generated:
            lang_text = str(lang)
            lang_tokens = cls._normalized_token_set(lang_text)
            for source in sources:
                if isinstance(source, dict):
                    source_text = " ".join(
                        str(source.get(field) or "") for field in ("language", "level")
                    )
                else:
                    source_text = str(source)
                if lang_tokens & cls._normalized_token_set(source_text):
                    sanitized.append(lang_text)
                    break

        return list(dict.fromkeys(sanitized))

    @classmethod
    def _sanitize_competencies(cls, competencies: list[str], source_text: str) -> list[str]:
        return [
            str(skill).strip()
            for skill in competencies
            if str(skill).strip() and cls._supported_by_source(str(skill), source_text)
        ]

    def _sanitize_generated_cv(
        self,
        cv: GeneratedCV,
        user: User,
        cv_analysis: dict | None,
        source_text: str,
    ) -> GeneratedCV:
        summary = self._remove_unsupported_numbered_sentences(cv.summary, source_text)
        if not summary and cv_analysis and isinstance(cv_analysis, dict):
            summary = str(cv_analysis.get("analysis") or "")

        return GeneratedCV(
            summary=summary,
            core_competencies=self._sanitize_competencies(cv.core_competencies, source_text),
            experience=self._sanitize_experience(cv.experience, user, cv_analysis, source_text),
            education=self._sanitize_education(cv.education, user, cv_analysis),
            certifications=cv.certifications,
            languages=self._sanitize_languages(cv.languages, user, cv_analysis),
        )

    def _build_job_text(self, job: Job) -> str:
        parts = [
            f"Titre du poste: {job.title}",
            f"Description: {job.description}",
        ]
        if job.company:
            parts.append(f"Entreprise: {job.company}")
        if job.city:
            parts.append(f"Ville: {job.city}")
        if job.salary_min or job.salary_max:
            parts.append(f"Salaire: {job.salary_min or ''} - {job.salary_max or ''}")
        if job.requirements:
            reqs = job.requirements
            if isinstance(reqs, dict):
                parts.append(f"Prérequis: {json.dumps(reqs, ensure_ascii=False)}")
            else:
                parts.append(f"Prérequis: {reqs}")
        return "\n".join(parts)

    def _call_llm(
        self, prompt: str, system_prompt: str, max_tokens: int = 2048, temperature: float = 0.2
    ) -> str:  # 0.2 = équilibre précision/créativité pour documents
        import time

        from openai import RateLimitError

        max_retries = 3
        base_delay = 2.0

        for attempt in range(max_retries):
            try:
                logger.info(
                    "[LLM] Appel au modèle %s (tentative %d/%d, max_tokens=%d)...",
                    settings.llm_model_id,
                    attempt + 1,
                    max_retries,
                    max_tokens,
                )
                response = self._client.chat.completions.create(
                    model=settings.llm_model_id,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=temperature,
                    max_tokens=max_tokens,
                )
                content = response.choices[0].message.content or ""
                logger.info(
                    "[LLM] Réponse reçue (%d caractères).",
                    len(content),
                )
                return content
            except RateLimitError:
                if attempt < max_retries - 1:
                    delay = base_delay * (2**attempt)
                    logger.warning(
                        f"RateLimitError. Retrying in {delay}s (attempt {attempt + 1}/{max_retries})..."
                    )
                    time.sleep(delay)
                else:
                    logger.exception("LLM call failed after retries due to RateLimitError")
                    raise DocumentGenerationError(
                        "Le service IA est très sollicité. Veuillez réessayer dans quelques minutes."
                    ) from None
            except Exception:
                logger.exception("LLM call failed during document generation")
                raise DocumentGenerationError(
                    "Le service IA est temporairement indisponible."
                ) from None
        return ""

    def _parse_json(self, content: str) -> dict | None:
        """Tente de parser le JSON retourne par le LLM.

        Gere les cas suivants :
        - JSON pur
        - JSON entoure de texte
        - JSON dans un bloc markdown ```json ... ```
        """
        if not content or not content.strip():
            return None

        json_str = self._extract_json_str(content)
        if not json_str:
            return None

        json_str = self._cleanup_json(json_str)
        try:
            return json.loads(json_str)
        except json.JSONDecodeError:
            pass
        return None

    def _extract_json_str(self, content: str) -> str | None:
        """Extrait la chaine JSON du contenu du LLM."""
        # JSON pur
        try:
            json.loads(content)
            return content
        except json.JSONDecodeError:
            pass

        # Bloc markdown ```json ... ``` ou ``` ... ```
        md_match = re.search(r"```(?:json)?\s*(\{.*\})\s*```", content, re.DOTALL)
        if md_match:
            return md_match.group(1)

        # JSON embarque dans du texte
        match = re.search(r"\{.*\}", content, re.DOTALL)
        if match:
            return match.group()

        return None

    @staticmethod
    def _cleanup_json(text: str) -> str:
        """Nettoie les erreurs JSON courantes des LLM."""
        text = re.sub(r",\s*([}\]])", r"\1", text)
        return text

    def generate_cv_content(
        self, user: User, job: Job, cv_analysis: dict | None = None
    ) -> GeneratedCV:
        """Génère le contenu structuré du CV via LLM, avec validation et régénération.

        Args:
            user: Profil utilisateur.
            job: Offre d'emploi cible.
            cv_analysis: Analyse du vrai CV uploadé (optionnel mais fortement recommandé).
        """
        logger.info("[CV] Construction du profil candidat...")
        profile_text = self._build_profile_text(user, cv_analysis=cv_analysis)
        source_text = self._source_text_from_analysis(cv_analysis, fallback=profile_text)
        job_text = self._build_job_text(job)
        prompt = CV_GENERATION_PROMPT.format(profile=profile_text, job=job_text)

        max_attempts = 2
        last_cv = None

        for attempt in range(max_attempts):
            logger.info(
                "[CV] Génération du contenu via LLM (tentative %d/%d)...", attempt + 1, max_attempts
            )
            content = self._call_llm(
                prompt,
                "Tu génères des CV ATS-optimisés en français. Réponds UNIQUEMENT avec du JSON valide.",
                max_tokens=2500,
            )

            logger.info("[CV] Parsing de la réponse JSON...")
            data = self._parse_json(content)
            if not data:
                logger.error(
                    "[CV] Echec du parsing JSON. Contenu brut du LLM (500 premiers caractères) : %r",
                    content[:500] if content else "(vide)",
                )
                if attempt == max_attempts - 1:
                    raise DocumentGenerationError("Impossible de générer le contenu du CV.")
                continue

            cv = GeneratedCV(
                summary=data.get("summary", ""),
                core_competencies=data.get("core_competencies", []),
                experience=data.get("experience", []),
                education=data.get("education", []),
                certifications=self._filter_certifications(
                    data.get("certifications", []), cv_analysis
                ),
                languages=data.get("languages", []),
            )
            cv = self._sanitize_generated_cv(cv, user, cv_analysis, source_text)

            warnings = self._validate_cv_content(cv, job)
            if warnings:
                logger.warning("[CV] Validation warnings (attempt %d): %s", attempt + 1, warnings)
                if attempt < max_attempts - 1:
                    correction_note = (
                        "\n\nATTENTION — Corrige ces problèmes dans ta réponse :\n"
                        + "\n".join(f"- {w}" for w in warnings)
                    )
                    prompt += correction_note
                    continue

            last_cv = cv
            break

        if last_cv is None:
            raise DocumentGenerationError("Impossible de générer un CV de qualité suffisante.")

        logger.info(
            "[CV] Contenu structuré OK — %d expérience(s), %d formation(s), %d compétence(s).",
            len(last_cv.experience),
            len(last_cv.education),
            len(last_cv.core_competencies),
        )
        return last_cv

    def generate_cover_letter_content(
        self,
        user: User,
        job: Job,
        match_analysis: str | None = None,
        cv_analysis: dict | None = None,
    ) -> GeneratedCoverLetter:
        """Génère le contenu de la lettre de motivation via LLM.

        Args:
            user: Profil utilisateur.
            job: Offre d'emploi cible.
            match_analysis: Analyse de correspondance profil/poste (optionnel).
            cv_analysis: Analyse du vrai CV uploadé (optionnel mais fortement recommandé).
        """
        logger.info("[Lettre] Construction du profil candidat...")
        profile_text = self._build_profile_text(user, cv_analysis=cv_analysis)
        source_text = self._source_text_from_analysis(cv_analysis, fallback=profile_text)
        job_text = self._build_job_text(job)
        prompt = COVER_LETTER_PROMPT.format(
            profile=profile_text,
            job=job_text,
            match_analysis=match_analysis or "Non disponible",
        )

        logger.info("[Lettre] Génération du contenu via LLM...")
        content = self._call_llm(
            prompt,
            "Tu rédiges des lettres de motivation professionnelles en français. Réponds UNIQUEMENT avec du JSON valide.",
        )
        logger.info("[Lettre] Parsing de la réponse JSON...")
        data = self._parse_json(content)
        if not data:
            logger.error(
                "[Lettre] Echec du parsing JSON. Contenu brut du LLM (500 premiers caractères) : %r",
                content[:500] if content else "(vide)",
            )
            raise DocumentGenerationError("Impossible de générer la lettre de motivation.")

        logger.info(
            "[Lettre] Contenu structuré OK — %d paragraphe(s).",
            len(data.get("body_paragraphs", [])),
        )
        body_paragraphs = [
            cleaned
            for paragraph in data.get("body_paragraphs", [])
            if (cleaned := self._remove_unsupported_numbered_sentences(str(paragraph), source_text))
        ]
        return GeneratedCoverLetter(
            recipient=data.get("recipient", "À l'attention du recruteur"),
            subject=data.get("subject", ""),
            body_paragraphs=body_paragraphs,
            closing=data.get("closing", "Sincères salutations"),
        )

    def _validate_cv_content(self, cv: GeneratedCV, job: Job) -> list[str]:
        """Vérifie que le CV respecte les standards de qualité.

        Returns:
            Liste des avertissements (vide = CV valide).
        """
        warnings: list[str] = []

        if len(cv.summary) < 100:
            warnings.append(f"Résumé trop court ({len(cv.summary)} caractères, minimum 100)")
        if len(cv.summary) > 600:
            warnings.append(f"Résumé trop long ({len(cv.summary)} caractères, maximum 600)")

        if len(cv.core_competencies) < 5:
            warnings.append(f"Trop peu de compétences ({len(cv.core_competencies)}, minimum 5)")
        if len(cv.core_competencies) > 14:
            warnings.append(f"Trop de compétences ({len(cv.core_competencies)}, maximum 14)")

        weak_verbs = {"géré", "réalisé", "effectué", "participé", "travaillé", "aidé"}
        for exp in cv.experience:
            for bullet in exp.get("bullets", []):
                first_word = bullet.split()[0].lower().rstrip(",.:") if bullet else ""
                if first_word in weak_verbs:
                    warnings.append(
                        f"Verbe faible détecté dans '{exp.get('company', '?')}': '{first_word}'"
                    )

        job_keywords = self._extract_job_keywords(job)
        cv_text = (cv.summary + " " + " ".join(cv.core_competencies)).lower()
        missing_keywords = [kw for kw in job_keywords[:5] if kw.lower() not in cv_text]
        if missing_keywords:
            warnings.append(
                f"Mots-clés importants de l'offre absents du CV : {', '.join(missing_keywords)}"
            )

        return warnings

    def _extract_job_keywords(self, job: Job) -> list[str]:
        """Extrait les mots-clés techniques d'une offre d'emploi."""
        text = f"{job.title} {job.description}"
        if job.requirements and isinstance(job.requirements, dict):
            text += " " + " ".join(str(v) for v in job.requirements.values())

        stop_words = {
            "le",
            "la",
            "les",
            "de",
            "du",
            "des",
            "un",
            "une",
            "et",
            "ou",
            "pour",
            "dans",
            "sur",
            "avec",
            "sans",
            "par",
            "est",
            "sont",
            "être",
            "avoir",
            "nous",
            "vous",
            "ils",
            "elle",
            "que",
            "qui",
            "quoi",
            "dont",
        }

        words = re.findall(r"\b[a-zA-ZÀ-ÿ]{3,}\b", text)
        freq: dict[str, int] = {}
        for word in words:
            w_lower = word.lower()
            if w_lower not in stop_words:
                freq[w_lower] = freq.get(w_lower, 0) + 1

        return sorted(freq, key=freq.get, reverse=True)[:10]

    def _set_cell_font(self, cell, name: str = "Calibri", size: int = 11, bold: bool = False):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = name
                run.font.size = Pt(size)
                run.font.bold = bold

    def build_cv_docx(self, cv: GeneratedCV, user: User, job: Job) -> BytesIO:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.15

        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(2)
        name_run = name_para.add_run(user.name.upper())
        name_run.font.name = "Calibri"
        name_run.font.size = Pt(22)
        name_run.bold = True
        name_run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)

        if job.title:
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(4)
            title_run = title_para.add_run(job.title)
            title_run.font.name = "Calibri"
            title_run.font.size = Pt(12)
            title_run.italic = True
            title_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        contact_parts = []
        if user.phone:
            contact_parts.append(f"{user.phone}")
        if user.email:
            contact_parts.append(f"{user.email}")
        location = ""
        if user.city:
            location = user.city
        if user.country and user.country not in (location or ""):
            location = f"{location}, {user.country}" if location else user.country
        if location:
            contact_parts.append(f"{location}")
        if user.linkedin_url:
            ln_short = user.linkedin_url.replace("https://", "").replace("www.", "").rstrip("/")
            contact_parts.append(f"{ln_short}")
        if user.github_url:
            gh_short = user.github_url.replace("https://", "").replace("www.", "").rstrip("/")
            contact_parts.append(f"{gh_short}")

        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(8)
            contact_run = contact_para.add_run("  ·  ".join(contact_parts))
            contact_run.font.name = "Calibri"
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        self._add_thick_separator(doc)

        self._add_section_heading(doc, "RÉSUMÉ PROFESSIONNEL")
        summary_para = doc.add_paragraph()
        summary_para.paragraph_format.space_after = Pt(8)
        summary_run = summary_para.add_run(cv.summary)
        summary_run.font.name = "Calibri"
        summary_run.font.size = Pt(10.5)
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if cv.core_competencies:
            self._add_section_heading(doc, "COMPÉTENCES CLÉS")
            self._add_skills_table(doc, cv.core_competencies)

        if cv.experience:
            self._add_section_heading(doc, "EXPÉRIENCES PROFESSIONNELLES")
            for i, exp in enumerate(cv.experience):
                self._add_experience_entry(doc, exp, is_last=(i == len(cv.experience) - 1))

        if cv.education:
            self._add_section_heading(doc, "FORMATION")
            for edu in cv.education:
                self._add_education_entry(doc, edu)

        certs_and_langs = []
        if cv.certifications:
            certs_and_langs.extend(cv.certifications)
        if cv.languages:
            certs_and_langs.extend(cv.languages)

        if certs_and_langs:
            self._add_section_heading(doc, "LANGUES & CERTIFICATIONS")
            for item in certs_and_langs:
                item_para = doc.add_paragraph(style="List Bullet")
                item_para.paragraph_format.space_after = Pt(2)
                item_run = item_para.add_run(item)
                item_run.font.name = "Calibri"
                item_run.font.size = Pt(10.5)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def build_cover_letter_docx(
        self, letter: GeneratedCoverLetter, user: User, job: Job
    ) -> BytesIO:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.1)
            section.right_margin = Inches(1.1)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.3

        sender_lines = [user.name]
        if user.address:
            sender_lines.append(user.address)
        if user.city or user.country:
            loc = ", ".join(filter(None, [user.city, user.country]))
            sender_lines.append(loc)
        if user.phone:
            sender_lines.append(user.phone)
        if user.email:
            sender_lines.append(user.email)

        sender_para = doc.add_paragraph()
        sender_para.paragraph_format.space_after = Pt(16)
        for i, line in enumerate(sender_lines):
            run = sender_para.add_run(line)
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
            if i == 0:
                run.bold = True
            if i < len(sender_lines) - 1:
                sender_para.add_run("\n")

        today = datetime.now().strftime("%d %B %Y")
        date_para = doc.add_paragraph(today)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.paragraph_format.space_after = Pt(16)
        for run in date_para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)

        recipient_para = doc.add_paragraph()
        recipient_para.paragraph_format.space_after = Pt(4)
        rec_run = recipient_para.add_run(letter.recipient)
        rec_run.font.name = "Calibri"
        rec_run.font.size = Pt(10.5)

        if job.company:
            company_para = doc.add_paragraph(job.company)
            company_para.paragraph_format.space_after = Pt(16)
            for run in company_para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)

        subject_para = doc.add_paragraph()
        subject_para.paragraph_format.space_after = Pt(16)
        subj_label = subject_para.add_run("Objet : ")
        subj_label.font.name = "Calibri"
        subj_label.font.size = Pt(10.5)
        subj_label.bold = True
        subj_content = subject_para.add_run(letter.subject)
        subj_content.font.name = "Calibri"
        subj_content.font.size = Pt(10.5)
        subj_content.bold = True
        subj_content.underline = True

        body = letter.body_paragraphs or []
        for para_text in body:
            para = doc.add_paragraph(para_text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Pt(24)
            para.paragraph_format.space_after = Pt(10)
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)

        doc.add_paragraph()
        closing_para = doc.add_paragraph(letter.closing)
        closing_para.paragraph_format.space_after = Pt(32)
        for run in closing_para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)

        sig_para = doc.add_paragraph()
        sig_name = sig_para.add_run(user.name)
        sig_name.font.name = "Calibri"
        sig_name.font.size = Pt(11)
        sig_name.bold = True

        contact_parts = []
        if user.phone:
            contact_parts.append(user.phone)
        if user.email:
            contact_parts.append(user.email)
        if contact_parts:
            sig_para.add_run("\n")
            sig_contact = sig_para.add_run("  ·  ".join(contact_parts))
            sig_contact.font.name = "Calibri"
            sig_contact.font.size = Pt(9.5)
            sig_contact.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def _add_section_heading(self, doc: Document, text: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
        fmt = para.paragraph_format
        fmt.space_before = Pt(10)
        fmt.space_after = Pt(4)

        ppr = para._p.get_or_add_pPr()
        pbdr = ppr.makeelement(qn("w:pBdr"), {})
        bottom = pbdr.makeelement(
            qn("w:bottom"),
            {
                qn("w:val"): "single",
                qn("w:sz"): "6",
                qn("w:space"): "1",
                qn("w:color"): "1A568E",
            },
        )
        pbdr.append(bottom)
        ppr.append(pbdr)

    def _add_thick_separator(self, doc: Document) -> None:
        """Ajoute une double ligne de séparation colorée."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(8)
        ppr = para._p.get_or_add_pPr()
        pbdr = ppr.makeelement(qn("w:pBdr"), {})
        bottom = pbdr.makeelement(
            qn("w:bottom"),
            {
                qn("w:val"): "double",
                qn("w:sz"): "6",
                qn("w:space"): "1",
                qn("w:color"): "1A568E",
            },
        )
        pbdr.append(bottom)
        ppr.append(pbdr)

    def _add_skills_table(self, doc: Document, skills: list[str]) -> None:
        """Affiche les compétences en tableau 3 colonnes pour économiser l'espace."""
        n_cols = 3
        n_rows = (len(skills) + n_cols - 1) // n_cols
        if n_rows == 0:
            n_rows = 1

        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.style = "Table Grid"

        for idx, skill in enumerate(skills):
            row_idx = idx // n_cols
            col_idx = idx % n_cols
            cell = table.rows[row_idx].cells[col_idx]
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            run = para.add_run(skill)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _add_experience_entry(self, doc: Document, exp: dict, is_last: bool = False) -> None:
        """Ajoute une entrée d'expérience avec mise en forme cohérente."""
        company = exp.get("company", "")
        location = exp.get("location", "")
        title = exp.get("title", "")
        start = exp.get("start_date", "")
        end = exp.get("end_date", "Présent")

        header_para = doc.add_paragraph()
        header_para.paragraph_format.space_before = Pt(6)
        header_para.paragraph_format.space_after = Pt(1)

        company_run = header_para.add_run(company.upper())
        company_run.font.name = "Calibri"
        company_run.font.size = Pt(11)
        company_run.bold = True
        company_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        if location:
            loc_run = header_para.add_run(f"  —  {location}")
            loc_run.font.name = "Calibri"
            loc_run.font.size = Pt(10)
            loc_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(3)

        title_run = title_para.add_run(title)
        title_run.font.name = "Calibri"
        title_run.font.size = Pt(10.5)
        title_run.italic = True
        title_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        if start or end:
            date_run = title_para.add_run(f"   [{start} – {end}]")
            date_run.font.name = "Calibri"
            date_run.font.size = Pt(9.5)
            date_run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
            date_run.bold = True

        for bullet in exp.get("bullets", []):
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.2)
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after = Pt(1)
            bullet_run = bp.add_run(bullet)
            bullet_run.font.name = "Calibri"
            bullet_run.font.size = Pt(10.5)

        if not is_last:
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(4)
            sep.paragraph_format.space_after = Pt(0)

    def _add_education_entry(self, doc: Document, edu: dict) -> None:
        """Ajoute une entrée de formation."""
        degree = edu.get("degree", "")
        institution = edu.get("institution", "")
        year = edu.get("year", "")

        edu_para = doc.add_paragraph()
        edu_para.paragraph_format.space_before = Pt(4)
        edu_para.paragraph_format.space_after = Pt(2)

        deg_run = edu_para.add_run(degree)
        deg_run.font.name = "Calibri"
        deg_run.font.size = Pt(10.5)
        deg_run.bold = True

        if institution:
            inst_run = edu_para.add_run(f"  —  {institution}")
            inst_run.font.name = "Calibri"
            inst_run.font.size = Pt(10.5)
            inst_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        if year and year.lower() not in ("", "en cours"):
            year_run = edu_para.add_run(f"  ({year})")
            year_run.font.name = "Calibri"
            year_run.font.size = Pt(9.5)
            year_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    @staticmethod
    def _pdf_multi_cell(pdf: FPDF, h: float, txt: str) -> None:
        """Safe multi_cell that always resets x to l_margin before rendering.

        Root cause of the crash: fpdf2's multi_cell() leaves self.x at the right
        edge of the last rendered fragment. The NEXT call with w=0 then computes
        available_width = page_width - r_margin - current_x ≈ 0, which raises
        "Not enough horizontal space to render a single character".

        Fix: unconditionally reset x to l_margin before every multi_cell call.
        """
        clean = (txt or "").strip()
        if not clean:
            return
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, h, clean)

    def build_cv_pdf(self, cv: GeneratedCV, user: User, job: Job) -> BytesIO:
        pdf = FPDF()
        pdf.add_page()

        margin = 15
        pdf.set_margins(margin, margin, margin)
        pdf.set_auto_page_break(auto=True, margin=15)

        if FONT_PATH and BOLD_FONT_PATH:
            pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
            pdf.add_font("DejaVu", "B", BOLD_FONT_PATH, uni=True)
            if ITALIC_FONT_PATH:
                pdf.add_font("DejaVu", "I", ITALIC_FONT_PATH, uni=True)
        else:
            logger.warning("DejaVu fonts not found, using built-in Helvetica")

        pdf.set_font(PDF_FONT, "B", 18)
        pdf.cell(0, 10, user.name, new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(2)

        contact_parts = []
        if user.phone:
            contact_parts.append(user.phone)
        if user.email:
            contact_parts.append(user.email)
        if user.city or user.address:
            loc = user.city or ""
            if user.address:
                loc = f"{user.address}, {loc}" if loc else user.address
            contact_parts.append(loc)
        if user.linkedin_url:
            contact_parts.append(f"LinkedIn: {user.linkedin_url}")
        if user.github_url:
            contact_parts.append(f"GitHub: {user.github_url}")

        pdf.set_font(PDF_FONT, "", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 6, " | ".join(contact_parts), new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(4)

        self._pdf_section(pdf, "Résumé Professionnel")
        pdf.set_font(PDF_FONT, "", 10)
        self._pdf_multi_cell(pdf, 5.5, cv.summary)
        pdf.ln(2)

        if cv.core_competencies:
            self._pdf_section(pdf, "Compétences Clés")
            pdf.set_font(PDF_FONT, "", 10)
            comps = ", ".join(cv.core_competencies)
            self._pdf_multi_cell(pdf, 5.5, comps)
            pdf.ln(2)

        if cv.experience:
            self._pdf_section(pdf, "Expérience Professionnelle")
            for exp in cv.experience:
                company = exp.get("company", "")
                location = exp.get("location", "")
                title = exp.get("title", "")
                start = exp.get("start_date", "")
                end = exp.get("end_date", "")

                pdf.set_font(PDF_FONT, "B", 11)
                line = company
                if location:
                    line += f"  |  {location}"
                pdf.cell(0, 6, line, new_x="LMARGIN", new_y="NEXT")

                pdf.set_font(PDF_FONT, "I", 10)
                line2 = title
                if start or end:
                    line2 += f"  —  {start} - {end}"
                pdf.cell(0, 5.5, line2, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)

                pdf.set_font(PDF_FONT, "", 10)
                for bullet in exp.get("bullets", []):
                    self._pdf_multi_cell(pdf, 5, f"- {bullet}")
                pdf.ln(2)

        if cv.education:
            self._pdf_section(pdf, "Formation")
            for edu in cv.education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                year = edu.get("year", "")

                pdf.set_font(PDF_FONT, "B", 10)
                line = degree
                if institution:
                    line += f"  —  {institution}"
                pdf.cell(0, 5.5, line, new_x="LMARGIN", new_y="NEXT")

                if year and year.lower() not in ("", "en cours"):
                    pdf.set_font(PDF_FONT, "", 9)
                    pdf.set_text_color(100, 100, 100)
                    pdf.cell(0, 5, f"({year})", new_x="LMARGIN", new_y="NEXT")
                    pdf.set_text_color(0, 0, 0)
                pdf.ln(1)

        certs_and_langs = []
        if cv.certifications:
            certs_and_langs.extend(cv.certifications)
        if cv.languages:
            certs_and_langs.extend(cv.languages)

        if certs_and_langs:
            self._pdf_section(pdf, "Langues & Certifications")
            pdf.set_font(PDF_FONT, "", 10)
            for item in certs_and_langs:
                self._pdf_multi_cell(pdf, 5.5, f"- {item}")
            pdf.ln(2)

        buf = BytesIO()
        pdf.output(buf)
        buf.seek(0)
        return buf

    def build_cover_letter_pdf(self, letter: GeneratedCoverLetter, user: User, job: Job) -> BytesIO:
        pdf = FPDF()
        pdf.add_page()

        margin = 15
        pdf.set_margins(margin, margin, margin)
        pdf.set_auto_page_break(auto=True, margin=15)

        if FONT_PATH and BOLD_FONT_PATH:
            pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
            pdf.add_font("DejaVu", "B", BOLD_FONT_PATH, uni=True)
            if ITALIC_FONT_PATH:
                pdf.add_font("DejaVu", "I", ITALIC_FONT_PATH, uni=True)
        else:
            logger.warning("DejaVu fonts not found, using built-in Helvetica")

        today = datetime.now().strftime("%d %B %Y")
        pdf.set_font(PDF_FONT, "", 10)
        pdf.cell(0, 6, today, new_x="LMARGIN", new_y="NEXT", align="R")
        pdf.ln(8)

        pdf.set_font(PDF_FONT, "B", 12)
        self._pdf_multi_cell(pdf, 6, letter.subject)
        pdf.ln(6)

        pdf.set_font(PDF_FONT, "", 11)
        for para_text in letter.body_paragraphs:
            self._pdf_multi_cell(pdf, 6, para_text)
            pdf.ln(3)

        pdf.ln(4)
        pdf.set_font(PDF_FONT, "", 11)
        self._pdf_multi_cell(pdf, 6, letter.closing)
        pdf.ln(8)

        pdf.set_font(PDF_FONT, "B", 11)
        pdf.cell(0, 6, user.name, new_x="LMARGIN", new_y="NEXT")

        contact_line = []
        if user.phone:
            contact_line.append(user.phone)
        if user.email:
            contact_line.append(user.email)

        if contact_line:
            pdf.set_font(PDF_FONT, "", 9)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(0, 5, " | ".join(contact_line), new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)

        buf = BytesIO()
        pdf.output(buf)
        buf.seek(0)
        return buf

    def _pdf_section(self, pdf: FPDF, title: str) -> None:
        pdf.set_font(PDF_FONT, "B", 12)
        pdf.set_text_color(26, 86, 142)
        pdf.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(26, 86, 142)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(3)
        pdf.set_text_color(0, 0, 0)

    def generate_cv(
        self,
        user: User,
        job: Job,
        fmt: str = "docx",
        cv_analysis: dict | None = None,
    ) -> tuple[BytesIO, str]:
        """Génère un CV ATS-optimisé.

        Args:
            user: Profil utilisateur.
            job: Offre d'emploi cible.
            fmt: Format de sortie ('docx' ou 'pdf').
            cv_analysis: Analyse du vrai CV uploadé — si fourni, le contenu généré
                         sera fidèle au vrai parcours de l'utilisateur.
        """
        logger.info(
            "[CV] Démarrage génération — user=%s (id=%s), poste='%s' (id=%s), format=%s%s",
            user.name,
            user.id,
            job.title,
            job.id,
            fmt,
            " [avec analyse CV]" if cv_analysis else " [sans CV uploadé]",
        )
        cv = self.generate_cv_content(user, job, cv_analysis=cv_analysis)
        name_clean = self._sanitize_filename(user.name)
        job_clean = self._sanitize_filename(job.title)
        filename = f"CV_{name_clean}_{job_clean}"[:80]

        if fmt == "pdf":
            logger.info("[CV] Rendu PDF en cours...")
            buf = self.build_cv_pdf(cv, user, job)
            logger.info(
                "[CV] PDF généré avec succès — fichier: %s.pdf (%d octets)",
                filename,
                buf.getbuffer().nbytes,
            )
            return buf, f"{filename}.pdf"

        logger.info("[CV] Rendu DOCX en cours...")
        buf = self.build_cv_docx(cv, user, job)
        logger.info(
            "[CV] DOCX généré avec succès — fichier: %s.docx (%d octets)",
            filename,
            buf.getbuffer().nbytes,
        )
        return buf, f"{filename}.docx"

    def generate_cover_letter(
        self,
        user: User,
        job: Job,
        fmt: str = "docx",
        match_analysis: str | None = None,
        cv_analysis: dict | None = None,
    ) -> tuple[BytesIO, str]:
        """Génère une lettre de motivation.

        Args:
            user: Profil utilisateur.
            job: Offre d'emploi cible.
            fmt: Format de sortie ('docx' ou 'pdf').
            match_analysis: Analyse de correspondance profil/poste.
            cv_analysis: Analyse du vrai CV uploadé.
        """
        logger.info(
            "[Lettre] Démarrage génération — user=%s (id=%s), poste='%s' (id=%s), format=%s%s",
            user.name,
            user.id,
            job.title,
            job.id,
            fmt,
            " [avec analyse CV]" if cv_analysis else " [sans CV uploadé]",
        )
        letter = self.generate_cover_letter_content(
            user, job, match_analysis, cv_analysis=cv_analysis
        )
        name_clean = self._sanitize_filename(user.name)
        job_clean = self._sanitize_filename(job.title)
        filename = f"Lettre_Motivation_{name_clean}_{job_clean}"[:80]

        if fmt == "pdf":
            logger.info("[Lettre] Rendu PDF en cours...")
            buf = self.build_cover_letter_pdf(letter, user, job)
            logger.info(
                "[Lettre] PDF généré avec succès — fichier: %s.pdf (%d octets)",
                filename,
                buf.getbuffer().nbytes,
            )
            return buf, f"{filename}.pdf"

        logger.info("[Lettre] Rendu DOCX en cours...")
        buf = self.build_cover_letter_docx(letter, user, job)
        logger.info(
            "[Lettre] DOCX généré avec succès — fichier: %s.docx (%d octets)",
            filename,
            buf.getbuffer().nbytes,
        )
        return buf, f"{filename}.docx"
